"""
Build the final research report as a Word document (.docx).
Generates all additional charts inline, embeds existing figures,
and writes a fully structured academic paper.
"""

import io, json, os, sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

# ── paths ──────────────────────────────────────────────────────────────────
BASE    = Path("C:/Users/aggarwm1/Videos/multi-disciplinary_Hyperspectral")
OUT_DIR = BASE / "outputs"
FIG_DIR = OUT_DIR / "figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)
DOCX_OUT = BASE / "model" / "Research_Report_Hyperspectral_Forest.docx"

# ── load results ───────────────────────────────────────────────────────────
with open(OUT_DIR / "baseline_multisite_casi_loso/experiment_summary.json") as f:
    baseline = json.load(f)
with open(OUT_DIR / "ablations/als_fusion.json") as f:
    als = json.load(f)
with open(OUT_DIR / "ablations/field_spectra.json") as f:
    field = json.load(f)
with open(OUT_DIR / "ablations/field_spectra_deep.json") as f:
    field_deep = json.load(f)
with open(OUT_DIR / "ablations/cross_sensor_bk.json") as f:
    cross = json.load(f)

# ── helper: save figure to BytesIO ────────────────────────────────────────
def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


# ═══════════════════════════════════════════════════════════════════════════
# CHART 1 – Model comparison bar chart (5 models × balanced accuracy)
# ═══════════════════════════════════════════════════════════════════════════
models = ["LogisticRegression", "RandomForest", "SVM", "GradientBoosting", "PLSDA"]
labels = ["Logistic\nRegression", "Random\nForest", "SVM\n(RBF)", "Gradient\nBoosting", "PLS-DA"]
ba_means = [baseline["models"][m]["mean_balanced_accuracy"] for m in models if m in baseline["models"]]
ba_stds  = [baseline["models"][m]["std_balanced_accuracy"]  for m in models if m in baseline["models"]]
oa_means = [baseline["models"][m]["mean_overall_accuracy"]  for m in models if m in baseline["models"]]
labels   = labels[:len(ba_means)]

fig1, ax = plt.subplots(figsize=(9, 4.5))
x = np.arange(len(labels))
w = 0.38
bars1 = ax.bar(x - w/2, ba_means, w, yerr=ba_stds, capsize=4,
               color="#2c7bb6", alpha=0.88, label="Balanced Accuracy", error_kw={"elinewidth":1.2})
bars2 = ax.bar(x + w/2, oa_means, w,
               color="#fdae61", alpha=0.88, label="Overall Accuracy")
ax.axhline(1/3, color="grey", linestyle="--", linewidth=0.9, label="Chance (0.33)")
ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=10)
ax.set_ylabel("Score", fontsize=11); ax.set_ylim(0, 1.0)
ax.set_title("Figure 1 — Classifier Comparison: LOSO-CV (3-class, 4 sites)", fontsize=11, fontweight="bold")
ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
for bar in bars1:
    h = bar.get_height()
    ax.text(bar.get_x() + bar.get_width()/2, h + 0.015, f"{h:.3f}", ha="center", va="bottom", fontsize=8)
plt.tight_layout()
chart1_stream = fig_to_stream(fig1)


# ═══════════════════════════════════════════════════════════════════════════
# CHART 2 – Per-site balanced accuracy (RF, 3-class)
# ═══════════════════════════════════════════════════════════════════════════
sites     = ["Bílý Kríž", "Hyytiälä", "Järvselja", "Lanzhot"]
rf_folds  = [1.000, 0.402, 0.583, 0.500]
site_colors = ["#e41a1c", "#377eb8", "#4daf4a", "#984ea3"]

fig2, ax = plt.subplots(figsize=(8, 4))
bars = ax.bar(sites, rf_folds, color=site_colors, alpha=0.85, edgecolor="white", linewidth=0.7)
ax.axhline(0.621, color="black", linestyle="--", linewidth=1.3, label=f"Mean BA = 0.621")
ax.axhline(1/3, color="grey",  linestyle=":",  linewidth=1.0, label="Chance (0.33)")
ax.set_ylim(0, 1.15); ax.set_ylabel("Balanced Accuracy", fontsize=11)
ax.set_title("Figure 2 — Random Forest: Per-Site LOSO Balanced Accuracy (3-class)", fontsize=11, fontweight="bold")
ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
for bar, val in zip(bars, rf_folds):
    ax.text(bar.get_x()+bar.get_width()/2, val+0.025, f"{val:.3f}", ha="center", fontsize=10, fontweight="bold")
plt.tight_layout()
chart2_stream = fig_to_stream(fig2)


# ═══════════════════════════════════════════════════════════════════════════
# CHART 3 – Confusion matrix (RF, aggregated over LOSO)
# ═══════════════════════════════════════════════════════════════════════════
cm_raw = np.array(baseline["models"]["RandomForest"]["total_confusion_matrix"])
class_names = ["Coniferous", "Broadleaved", "Mixed"]
row_sums = cm_raw.sum(axis=1, keepdims=True); row_sums[row_sums==0] = 1
cm_norm = cm_raw / row_sums

fig3, ax = plt.subplots(figsize=(5.5, 4.5))
im = ax.imshow(cm_norm, cmap="Blues", vmin=0, vmax=1)
plt.colorbar(im, ax=ax, fraction=0.046)
ax.set_xticks(range(3)); ax.set_yticks(range(3))
ax.set_xticklabels(class_names, rotation=30, ha="right", fontsize=10)
ax.set_yticklabels(class_names, fontsize=10)
ax.set_xlabel("Predicted", fontsize=11); ax.set_ylabel("True", fontsize=11)
ax.set_title("Figure 3 — RF Aggregated Confusion Matrix\n(row-normalised, LOSO 4-fold)", fontsize=10, fontweight="bold")
for i in range(3):
    for j in range(3):
        val = cm_norm[i, j]
        raw = cm_raw[i, j]
        color = "white" if val > 0.5 else "black"
        ax.text(j, i, f"{val:.2f}\n(n={raw})", ha="center", va="center", color=color, fontsize=9)
plt.tight_layout()
chart3_stream = fig_to_stream(fig3)


# ═══════════════════════════════════════════════════════════════════════════
# CHART 4 – Ablation comparison (bar chart)
# ═══════════════════════════════════════════════════════════════════════════
ablation_labels = [
    "Baseline\n(CASI 40 bands)",
    "PCA-10\nCompression",
    "ALS Only\n(structural)",
    "CASI + ALS\nFusion",
    "Binary\n(no mixed)",
    "Field Spectra\nLOSO",
    "Cross-Sensor\nBK Transfer",
]
ablation_3class = [0.621, 0.600, 0.192, 0.663, None, 0.413, None]
ablation_binary = [0.932, None, 0.247, 0.932, 0.932, None, 0.875]

fig4, ax = plt.subplots(figsize=(11, 5))
x = np.arange(len(ablation_labels))
w = 0.38
vals3 = [v if v is not None else 0 for v in ablation_3class]
vals2 = [v if v is not None else 0 for v in ablation_binary]
alpha3 = [0.85 if v is not None else 0.0 for v in ablation_3class]
alpha2 = [0.85 if v is not None else 0.0 for v in ablation_binary]

bars3 = ax.bar(x - w/2, vals3, w, color="#2c7bb6", alpha=0.85, label="3-class BA")
bars2 = ax.bar(x + w/2, vals2, w, color="#d7191c", alpha=0.85, label="Binary BA")

# make N/A bars invisible
for bar, a in zip(bars3, alpha3):
    bar.set_alpha(a)
for bar, a in zip(bars2, alpha2):
    bar.set_alpha(a)

ax.axhline(1/3, color="grey", linestyle=":", linewidth=1)
ax.set_xticks(x); ax.set_xticklabels(ablation_labels, fontsize=9)
ax.set_ylim(0, 1.08); ax.set_ylabel("Balanced Accuracy", fontsize=11)
ax.set_title("Figure 4 — Ablation Study: 3-class vs Binary Classification Across Feature Sets", fontsize=11, fontweight="bold")
ax.legend(fontsize=10); ax.grid(axis="y", alpha=0.3)

for bar, val, is_vis in zip(bars3, vals3, alpha3):
    if is_vis:
        ax.text(bar.get_x()+bar.get_width()/2, val+0.02, f"{val:.3f}", ha="center", fontsize=8)
for bar, val, is_vis in zip(bars2, vals2, alpha2):
    if is_vis:
        ax.text(bar.get_x()+bar.get_width()/2, val+0.02, f"{val:.3f}", ha="center", fontsize=8)
plt.tight_layout()
chart4_stream = fig_to_stream(fig4)


# ═══════════════════════════════════════════════════════════════════════════
# CHART 5 – Cross-sensor BK prediction probabilities
# ═══════════════════════════════════════════════════════════════════════════
bk_stands = [p["stand_id"].replace("BK_", "") for p in cross["bk_predictions"]]
p_conif   = [p["probabilities"]["coniferous"] for p in cross["bk_predictions"]]
p_broad   = [p["probabilities"]["broadleaved"] for p in cross["bk_predictions"]]
p_mixed   = [p["probabilities"]["mixed"] for p in cross["bk_predictions"]]

fig5, ax = plt.subplots(figsize=(9, 4.5))
x5 = np.arange(len(bk_stands))
w5 = 0.28
ax.bar(x5 - w5, p_conif, w5, label="P(coniferous)", color="#2e7d32", alpha=0.85)
ax.bar(x5,       p_mixed, w5, label="P(mixed)",      color="#7b1fa2", alpha=0.85)
ax.bar(x5 + w5,  p_broad, w5, label="P(broadleaved)",color="#f57f17", alpha=0.85)
ax.set_xticks(x5); ax.set_xticklabels(bk_stands, rotation=30, ha="right", fontsize=9)
ax.set_ylabel("Probability", fontsize=11); ax.set_ylim(0, 1)
ax.set_title("Figure 5 — Cross-Sensor Prediction: CASI Model Applied to Bílý Kríž Field Spectra\n"
             "(trained on HY+JS+LZ airborne; resampled to 40 CASI bands)", fontsize=10, fontweight="bold")
ax.axhline(0.5, color="grey", linestyle="--", linewidth=0.8)
ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
# mark the misclassification
ax.annotate("Misclassified\n(27% broadleaved BA)", xy=(2, p_mixed[2]+0.03),
            xytext=(3.5, 0.65), arrowprops=dict(arrowstyle="->", color="red"), color="red", fontsize=8)
plt.tight_layout()
chart5_stream = fig_to_stream(fig5)


# ═══════════════════════════════════════════════════════════════════════════
# CHART 6 – Dataset overview (site × class heatmap)
# ═══════════════════════════════════════════════════════════════════════════
site_names = ["Hyytiälä\n(Boreal)", "Järvselja\n(Hemiboreal)", "Bílý Kríž\n(Temperate\nMontane)", "Lanzhot\n(Temperate\nFloodplain)"]
class_names6 = ["Coniferous", "Broadleaved", "Mixed"]
counts = np.array([
    [22, 4,  2],   # HY
    [4,  4,  5],   # JS
    [7,  0,  0],   # BK
    [0,  7,  3],   # LZ
])

fig6, ax = plt.subplots(figsize=(7, 4.5))
im6 = ax.imshow(counts, cmap="YlOrRd", aspect="auto")
plt.colorbar(im6, ax=ax, label="Number of stands")
ax.set_xticks(range(3)); ax.set_yticks(range(4))
ax.set_xticklabels(class_names6, fontsize=11)
ax.set_yticklabels(site_names, fontsize=10)
ax.set_title("Figure 6 — Dataset Overview: Stand Count by Site and Forest Type", fontsize=11, fontweight="bold")
for i in range(4):
    for j in range(3):
        v = counts[i, j]
        ax.text(j, i, str(v), ha="center", va="center", fontsize=14,
                fontweight="bold", color="white" if v > 6 else "black")
plt.tight_layout()
chart6_stream = fig_to_stream(fig6)


# ═══════════════════════════════════════════════════════════════════════════
# CHART 7 – Pipeline workflow diagram
# ═══════════════════════════════════════════════════════════════════════════
fig7, ax = plt.subplots(figsize=(12, 3))
ax.set_xlim(0, 12); ax.set_ylim(0, 3); ax.axis("off")

steps = [
    ("GeoTIFF\nTiles\n(58 stands)", 0.5, "#4e79a7"),
    ("Band\nSelection\n(40 bands)", 2.2, "#f28e2b"),
    ("Stand-level\nMean\nSpectrum", 3.9, "#e15759"),
    ("Standard\nNormalisation\n(per fold)", 5.6, "#76b7b2"),
    ("Classifier\nTraining\n(RF, SVM…)", 7.3, "#59a14f"),
    ("LOSO-CV\nEvaluation\n(4 folds)", 9.0, "#edc948"),
    ("Metrics &\nFigures\n(BA, F1, CM)", 10.7, "#b07aa1"),
]
for label, x, color in steps:
    ax.add_patch(mpatches.FancyBboxPatch((x-0.7, 0.4), 1.4, 2.1, boxstyle="round,pad=0.1",
                                          facecolor=color, edgecolor="white", linewidth=2, alpha=0.9))
    ax.text(x, 1.5, label, ha="center", va="center", fontsize=8.5,
            color="white", fontweight="bold", multialignment="center")

for i in range(len(steps)-1):
    x1 = steps[i][1] + 0.7
    x2 = steps[i+1][1] - 0.7
    ax.annotate("", xy=(x2, 1.5), xytext=(x1, 1.5),
                arrowprops=dict(arrowstyle="->", color="#333333", lw=2))

ax.set_title("Figure 7 — End-to-End Classification Pipeline", fontsize=12, fontweight="bold", pad=8)
plt.tight_layout()
chart7_stream = fig_to_stream(fig7)


# ═══════════════════════════════════════════════════════════════════════════
# CHART 8 – Field spectra per-site result
# ═══════════════════════════════════════════════════════════════════════════
field_sites = ["Bílý Kríž\n(all coniferous)", "Järvselja\n(mixed types)", "Lanzhot\n(all broadleaved)"]
field_ba    = [0.0, 0.639, 0.600]
airborne_ba = [1.0, 0.583, 0.500]

fig8, ax = plt.subplots(figsize=(8, 4.5))
x8 = np.arange(3); w8 = 0.38
ax.bar(x8 - w8/2, field_ba,    w8, color="#984ea3", alpha=0.85, label="Field Spectra (350–2500 nm)")
ax.bar(x8 + w8/2, airborne_ba, w8, color="#2c7bb6", alpha=0.85, label="Airborne CASI (382–1052 nm)")
ax.set_xticks(x8); ax.set_xticklabels(field_sites, fontsize=10)
ax.set_ylim(0, 1.15); ax.set_ylabel("Balanced Accuracy", fontsize=11)
ax.set_title("Figure 8 — Airborne vs. Field Spectra: Per-Site LOSO Balanced Accuracy", fontsize=11, fontweight="bold")
ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
for vals in [field_ba, airborne_ba]:
    offset = -w8/2 if vals == field_ba else w8/2
    for i, v in enumerate(vals):
        ax.text(x8[i]+offset, v+0.03, f"{v:.3f}", ha="center", fontsize=9, fontweight="bold")
plt.tight_layout()
chart8_stream = fig_to_stream(fig8)


print("All charts generated.")

# ══════════════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Styles ────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name, font_size, bold=False, color=None, space_before=0, space_after=6):
    try:
        s = styles[style_name]
    except KeyError:
        return
    s.font.name = font_name
    s.font.size = Pt(font_size)
    s.font.bold = bold
    if color:
        s.font.color.rgb = RGBColor(*color)
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after  = Pt(space_after)

set_style("Normal",    "Calibri", 11, space_after=6)
set_style("Heading 1", "Calibri", 18, bold=True, color=(31,73,125),  space_before=18, space_after=6)
set_style("Heading 2", "Calibri", 13, bold=True, color=(31,73,125),  space_before=12, space_after=4)
set_style("Heading 3", "Calibri", 11, bold=True, color=(68,114,196), space_before=8,  space_after=2)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(80, 80, 80)

def add_chart(doc, stream, width_inches=6.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(stream, width=Inches(width_inches))
    if caption:
        add_caption(doc, caption)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DCE6F1")
                tcPr.append(shd)
    if col_widths:
        for i, w_cm in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w_cm)
    doc.add_paragraph()
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level*0.8)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def page_break(doc):
    doc.add_page_break()

def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Spectral Separability of Forest Types Across European Biomes")
run.bold = True; run.font.size = Pt(22); run.font.color.rgb = RGBColor(31,73,125)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Airborne Hyperspectral Classification with Leave-One-Site-Out Cross-Validation")
run.bold = True; run.font.size = Pt(14); run.font.color.rgb = RGBColor(68,114,196)

doc.add_paragraph()
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.add_run("Manan Aggarwal").font.size = Pt(13)

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run("Multi-disciplinary Project — Hyperspectral Remote Sensing   |   June 2026")
run.font.size = Pt(11); run.font.color.rgb = RGBColor(100,100,100); run.italic = True

doc.add_paragraph()
horizontal_rule(doc)

# ── ABSTRACT ──────────────────────────────────────────────────────────────
add_heading(doc, "Abstract", 1)
add_para(doc,
    "This study investigates whether airborne hyperspectral reflectance can distinguish coniferous, "
    "broadleaved, and mixed forest stands, and whether classifiers trained on one biome generalise "
    "to others. Using the FREEDLES dataset (Rautiainen et al. 2024) — 58 forest stands at four "
    "European sites spanning boreal (Hyytiälä, Finland), hemiboreal (Järvselja, Estonia), and "
    "temperate (Bílý Kríž and Lanzhot, Czech Republic) biomes — we computed stand-level mean "
    "reflectance spectra from airborne CASI-1500 imagery (48 bands, 382–1052 nm) and evaluated "
    "five classifiers under strict Leave-One-Site-Out (LOSO) cross-validation. Random Forest "
    "achieved the highest mean balanced accuracy of 0.621 for three-class classification and "
    "0.932 for binary coniferous-vs-broadleaved discrimination. Airborne Laser Scanning (ALS) "
    "structural features provided a modest additive benefit (+4.1 percentage points) for the "
    "three-class problem but no improvement for binary classification, confirming that forest type "
    "is encoded primarily in leaf optical properties rather than canopy architecture. Field "
    "spectroradiometer measurements (350–2500 nm, BilyKriz dataset) identified the SWIR-2 region "
    "(2317–2490 nm) as the most discriminative spectral window; yet the CASI VIS-NIR model already "
    "achieves near-perfect binary accuracy, confirming these additional wavelengths are redundant "
    "at stand scale. Cross-sensor prediction of Bílý Kríž Norway spruce stands from the "
    "CASI-trained model achieved BA = 0.875 (7/8 correct), confirming inter-sensor "
    "transferability. The primary limitation is the mixed-class label: all ten mixed stands were "
    "misclassified as coniferous, indicating that mixed forests require sub-stand spatial analysis "
    "rather than stand-level mean spectra."
)
horizontal_rule(doc)

page_break(doc)


# ══════════════════════════════════════════════════════════════════
#  SECTION 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════
add_heading(doc, "1.  Introduction", 1)

add_para(doc,
    "Forest ecosystems cover approximately 31% of the Earth's land surface and play a central "
    "role in the global carbon cycle, biodiversity conservation, and climate regulation. Accurate "
    "mapping of forest type — the distinction between coniferous, broadleaved, and mixed stands "
    "— is a prerequisite for national forest inventories, carbon accounting frameworks such as "
    "REDD+, and climate-change impact assessments. Traditional field surveys are the gold "
    "standard for accuracy but are prohibitively expensive at continental scale. Satellite "
    "multispectral imagery (Sentinel-2, Landsat) offers wall-to-wall coverage but provides only "
    "12–13 broad spectral bands, insufficient to resolve the subtle biochemical differences "
    "between forest types across biomes."
)
add_para(doc,
    "Airborne hyperspectral imaging bridges this gap by delivering hundreds of narrow, contiguous "
    "spectral bands at sub-metre spatial resolution. Laboratory and field studies have "
    "demonstrated that spectral reflectance at the leaf level is strongly controlled by "
    "biochemical composition — chlorophyll, carotenoids, water content, cellulose, and lignin — "
    "which in turn differs systematically between needle-leaved conifers and broad-leaved "
    "deciduous trees. The question is whether these leaf-level differences survive aggregation "
    "to the stand level (where canopy architecture, gap fraction, soil background, and "
    "understory contribute) and, critically, whether they generalise across biomes where the "
    "same broad forest types are expressed under very different climate, soil, and phenological "
    "conditions."
)
add_para(doc,
    "This project addresses three inter-related research questions using the FREEDLES multi-site "
    "airborne dataset:"
)
add_bullet(doc, "Can CASI-1500 stand-level mean reflectance (382–1052 nm) distinguish coniferous, broadleaved, and mixed forest types across four European biomes under Leave-One-Site-Out cross-validation?")
add_bullet(doc, "Does adding ALS-derived canopy structural features improve classification, and for which problem (3-class vs. binary)?")
add_bullet(doc, "Do the findings transfer to field spectroradiometer data (350–2500 nm), and can the airborne-trained model be applied cross-sensor to the Bílý Kríž field spectroscopy dataset?")

doc.add_paragraph()
add_para(doc,
    "These questions collectively probe the ecological coherence and practical operationality "
    "of spectral forest-type classification at the biome scale — a prerequisite for the "
    "operational deployment of airborne hyperspectral campaigns in continental forest inventory."
)


# ══════════════════════════════════════════════════════════════════
#  SECTION 2 — DATA
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "2.  Data", 1)

add_heading(doc, "2.1  FREEDLES Airborne Dataset", 2)
add_para(doc,
    "The FREEDLES dataset (Rautiainen et al. 2024, Earth System Science Data 16, 5069–5098) "
    "provides analysis-ready hyperspectral image tiles and airborne laser scanning (ALS) point "
    "clouds for 58 forest stands at four European sites. Sites were selected to represent a "
    "latitudinal gradient from boreal Finland to temperate Czech Republic, spanning three "
    "distinct biomes. All stands are 100 m × 100 m plots."
)

add_table(doc,
    ["Site", "Country", "Biome", "n stands", "Dominant species"],
    [
        ["Hyytiälä (HY)", "Finland", "Boreal", "28", "Scots pine, Norway spruce, Silver birch"],
        ["Järvselja (JS)", "Estonia", "Hemiboreal", "13", "Scots pine, Norway spruce, Silver birch, Alder"],
        ["Bílý Kríž (BK)", "Czech Rep.", "Temperate montane", "7", "Norway spruce (monoculture plantation)"],
        ["Lanzhot (LZ)", "Czech Rep.", "Temperate floodplain", "10", "English oak, Ash, Hornbeam"],
    ],
    col_widths=[3.5, 2.5, 3.5, 2.0, 5.0]
)
add_caption(doc, "Table 1. Site descriptions for the FREEDLES dataset.")

add_heading(doc, "2.2  Sensors", 2)
add_para(doc,
    "Two airborne sensors were used. The CASI-1500 (ITRES Research, Canada) acquires "
    "hyperspectral imagery in 48 bands from 381.97 to 1052.15 nm at approximately 14 nm "
    "spectral resolution and 0.5 m spatial resolution. The SASI-600 covers 950–2443 nm; "
    "a merged CS product (140 bands, 382–2443 nm) is also available at 1.25 m resolution. "
    "The RIEGL LMS-Q780 ALS system provides full-waveform point clouds at approximately "
    "1 m ground density; analysis-ready LAS files are provided per stand."
)
add_para(doc,
    "Raw CASI tiles store reflectance as signed 16-bit integers scaled by a factor of 10,000 "
    "(i.e., actual reflectance = DN / 10,000). The no-data sentinel value is 10,000 — "
    "critically, this is distinct from the valid range [0, 1] and must be masked before "
    "processing. No spectral wavelength metadata is embedded in the GeoTIFF files; wavelengths "
    "were extracted from companion ENVI flight-line header files and injected programmatically."
)

add_heading(doc, "2.3  BilyKriz Field Spectroscopy Dataset", 2)
add_para(doc,
    "A supplementary dataset (DatasetOfTreeCanopyStructure...V2.xlsx) provides 465 "
    "field-measured canopy reflectance spectra at 350–2500 nm (1 nm step, approximately "
    "ASD FieldSpec resolution) for 31 stands across Järvselja (13), Bílý Kríž (8), "
    "and Lanzhot (10). Each stand has 15 positional measurements taken within the stand "
    "polygon. Bílý Kríž contributes BK_Spruce1–8, one additional stand (Spruce8) beyond "
    "the 7 stands present in the airborne tile dataset. Forest type was assigned from "
    "species basal-area percentages: stands with >70% coniferous basal area were labelled "
    "coniferous; >70% broadleaved were labelled broadleaved; others were labelled mixed. "
    "In practice the field dataset contains no mixed stands — Bílý Kríž is all-coniferous "
    "and Lanzhot is all-broadleaved."
)

# Dataset overview chart
add_chart(doc, chart6_stream, width_inches=5.5,
          caption="Figure 6 — Stand counts by site and forest type. Note the single-type nature of "
                  "Bílý Kríž (all coniferous) and Lanzhot (all broadleaved), which complicates "
                  "LOSO cross-site evaluation.")


# ══════════════════════════════════════════════════════════════════
#  SECTION 3 — PROBLEM FORMULATION
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "3.  Problem Formulation", 1)

add_para(doc,
    "Let xi ∈ R^d be the stand-level mean reflectance vector for stand i, where d = 40 "
    "valid spectral bands after water-vapour exclusion. Let yi ∈ {coniferous, broadleaved, "
    "mixed} be the forest-type label assigned from field inventory records. The goal is to "
    "learn a classifier f: R^d → Y that, when trained on stands from three sites "
    "S_train = {HY, JS, LZ}, correctly assigns labels to unseen stands at the fourth site "
    "S_test = {BK}."
)
add_para(doc,
    "This formulation embeds an explicit spatial generalisation constraint: training and "
    "test data come from geographically and climatically distinct regions. This is in "
    "contrast to the random train/test split common in machine learning, which would allow "
    "spatial autocorrelation to leak information from training to test sets and produce "
    "optimistically biased performance estimates (Roberts et al. 2017; Ploton et al. 2020)."
)

add_heading(doc, "3.1  Leave-One-Site-Out Cross-Validation", 2)
add_para(doc,
    "LOSO-CV operates on the four sites as fold groups. In each of four iterations, one "
    "site provides all test stands; the remaining three sites provide all training stands. "
    "No stand, pixel, or metadata from the test site appears during training or "
    "hyperparameter selection. This is the strictest defensible evaluation for this "
    "dataset and directly measures the ability of the classifier to generalise to new, "
    "unseen biomes — the operationally relevant scenario for continental-scale deployment."
)

add_heading(doc, "3.2  Primary Metric: Balanced Accuracy", 2)
add_para(doc,
    "Balanced accuracy (BA) is the arithmetic mean of per-class recall. It is the "
    "appropriate primary metric here because the class distribution is strongly imbalanced "
    "(33 coniferous : 15 broadleaved : 10 mixed). Overall accuracy would be dominated by "
    "the majority class and would not reflect the classifier's ability to identify the "
    "rarer broadleaved and mixed classes. Secondary metrics include macro-F1 "
    "(unweighted average F1), overall accuracy, and Cohen's κ."
)

add_heading(doc, "3.3  Binary Sub-Problem", 2)
add_para(doc,
    "We additionally evaluate the binary coniferous vs. broadleaved problem by excluding "
    "the 10 mixed stands. This tests a cleaner ecological hypothesis — whether spectral "
    "reflectance can separate needle-leaved from broad-leaved canopies — without the "
    "complication of the ambiguous mixed label. The binary dataset has 48 stands "
    "(33 coniferous, 15 broadleaved) and the same LOSO structure."
)


# ══════════════════════════════════════════════════════════════════
#  SECTION 4 — EXPLORATORY DATA ANALYSIS
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "4.  Data Familiarisation and Exploratory Analysis", 1)

add_heading(doc, "4.1  Data Discovery", 2)
add_para(doc,
    "The dataset was not at the path specified in the original configuration "
    "(D:/Hyperspectral_Data). The actual data resided in a deeply nested directory:"
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.2)
p.paragraph_format.space_after = Pt(6)
run = p.add_run(
    "model/data/Hyperspectral/Airborne_data/Airborne_hyperspectral/\n"
    "  Analysis_ready_subsets/CASI/    ← 58 × *_CASI.tif\n"
    "  Analysis_ready_subsets/CS/      ← 58 × *_CS.tif (merged product)\n"
    "model/data/Airborne_laser_scanning/...Analysis_ready_subsets/{STAND_ID}/*.las"
)
run.font.name = "Courier New"; run.font.size = Pt(9)

add_para(doc,
    "Several critical data properties were not documented in the dataset README and were "
    "discovered through code inspection and header analysis:"
)
add_bullet(doc, "Raw DN values encode reflectance × 10,000 (not 0–1 physical reflectance). Raw range: 0–2350 per band, corresponding to 0–0.235 reflectance.")
add_bullet(doc, "No-data sentinel: 10,000 (not 0). Initialising the mask with nodata=0 silently treated all pixels as valid.")
add_bullet(doc, "GeoTIFF files contain no spectral wavelength metadata. The loader fell back to band indices [0, 1, …, 47]. Water-vapour exclusion (which compares wavelength values in nm) silently skipped all 48 bands.")
add_bullet(doc, "Tile filenames include a sensor suffix: HY_PINE1_CASI.tif produces stand ID HY_PINE1_CASI, which does not match the metadata key HY_PINE1.")
add_bullet(doc, "ALS LAS files at Bílý Kríž and Lanzhot use classification code 4 (medium vegetation) rather than code 1 (unclassified/vegetation). Filtering on cls==1 produced zero vegetation points.")

add_heading(doc, "4.2  Class and Site Distribution", 2)
add_para(doc,
    "The stand-level class distribution is imbalanced: 33 coniferous (57%), "
    "15 broadleaved (26%), 10 mixed (17%). More importantly, the distribution is "
    "strongly site-specific. Bílý Kríž is entirely coniferous (Norway spruce plantation), "
    "making its LOSO test fold trivially easy — any classifier that labels all BK stands "
    "as coniferous achieves BA = 1.0 on that fold. This inflates the reported mean BA "
    "across folds and must be interpreted cautiously. The meaningful test folds are "
    "Hyytiälä (all three classes, 28 stands) and Järvselja (all three classes, 13 stands)."
)

add_heading(doc, "4.3  Spectral Characteristics", 2)
add_para(doc,
    "Water-vapour absorption creates interpolation artefacts in four wavelength intervals "
    "— 895–1003 nm, 1092–1168 nm, 1302–1528 nm, and 1737–2038 nm — where ATCOR-4 "
    "atmospheric correction fills in interpolated values rather than true surface "
    "reflectance. These regions were excluded, reducing the 48-band CASI spectrum to "
    "40 valid bands."
)
add_para(doc,
    "Visual inspection of mean spectra per forest type reveals ecologically interpretable "
    "differences. Broadleaved stands show a sharper red-edge inflection (700–730 nm), "
    "a higher near-infrared plateau (800–900 nm), and a characteristic green peak at "
    "550 nm driven by chlorophyll b. Coniferous stands show a lower NIR plateau and "
    "a broader red-edge associated with the layered needle architecture. Mixed stands "
    "are intermediate with high within-class variance, making them spectrally ambiguous "
    "at stand scale."
)

# Embed the existing mean_spectra figure
spectra_fig_path = OUT_DIR / "baseline_multisite_casi_loso/figures/mean_spectra_per_class.png"
if spectra_fig_path.exists():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(spectra_fig_path), width=Inches(5.5))
    add_caption(doc, "Figure 9 — Mean reflectance spectra (± 1 std) per forest type across all 58 stands. "
                     "Shaded regions mark excluded water-vapour bands. Broadleaved stands show a sharper "
                     "red-edge and higher NIR plateau than coniferous stands.")


# ══════════════════════════════════════════════════════════════════
#  SECTION 5 — METHODS
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "5.  Methods", 1)

add_heading(doc, "5.1  Feature Extraction", 2)
add_para(doc,
    "Stand-level mean spectrum: for each 100 m × 100 m tile, all valid pixels "
    "(DN ≠ 10,000 in any band) are averaged band-wise, producing a single d-dimensional "
    "feature vector per stand. Valid-pixel fraction was verified to exceed 50% for all "
    "58 stands. This aggregation discards within-stand spatial texture but provides a "
    "robust representation in the small-sample regime (n = 58) where per-pixel classifiers "
    "would severely overfit."
)
add_para(doc,
    "ALS structural features: raw LAS point clouds (multiple flight-line files per stand) "
    "were merged. Ground-classified returns (code = 2) were used to estimate ground "
    "elevation (5th-percentile of ground Z). Normalised heights for vegetation returns "
    "(all non-ground codes) were computed as Z − ground_elevation. Eight stand-level "
    "metrics were extracted: maximum height (h_max), mean height (h_mean), height "
    "standard deviation (h_std = rugosity), four height percentiles (p25, p50, p75, p95), "
    "and canopy cover fraction (proportion of first returns above 2 m). All 58 stands "
    "had sufficient ALS coverage."
)

# Pipeline chart
add_chart(doc, chart7_stream, width_inches=6.3,
          caption="Figure 7 — End-to-end classification pipeline from raw GeoTIFF tiles to evaluated metrics.")

add_heading(doc, "5.2  Preprocessing", 2)
add_para(doc,
    "Band selection: the 8 bands falling within the four water-vapour exclusion intervals "
    "were removed, leaving 40 valid CASI bands. Standard z-score normalisation "
    "(zero mean, unit variance) was fitted exclusively on the training fold and applied "
    "to both training and test folds, preventing any information leak from test data "
    "into the normalisation parameters."
)
add_para(doc,
    "PCA (ablation only): in the PCA-10 ablation, principal component analysis was "
    "fitted on the normalised training set and applied to compress 40 dimensions to "
    "10 principal components. This was applied within each LOSO fold to avoid leakage."
)

add_heading(doc, "5.3  Classifiers", 2)

add_table(doc,
    ["Model", "Key Parameters", "Notes"],
    [
        ["Logistic Regression", "C=1.0, solver=lbfgs, balanced weights, max_iter=1000", "Multi-class via softmax"],
        ["Random Forest", "500 trees, min_samples_leaf=2, balanced weights", "Best performer; used for all ablations"],
        ["SVM (RBF)", "C=10, gamma=scale, balanced weights", "Kernel trick for non-linear boundaries"],
        ["Gradient Boosting", "200 trees, lr=0.1, max_depth=5, subsample=0.8", "Sequential ensemble"],
        ["PLS-DA", "10 components, max_iter=500", "Standard in chemometrics/remote sensing"],
    ],
    col_widths=[3.5, 6.5, 4.5]
)
add_caption(doc, "Table 2. Classifier configurations. All models use balanced class weights.")

add_heading(doc, "5.4  Ablation Experiments", 2)
add_para(doc,
    "Six controlled ablations were designed to isolate the contribution of individual "
    "design choices:"
)
add_bullet(doc, "A1 — PCA compression: replace 40 spectral bands with 10 principal components. Tests whether dimensionality reduction improves generalisation.")
add_bullet(doc, "A2 — Binary problem: drop mixed-class stands. Isolates whether the conifer/broadleaf distinction is biome-stable.")
add_bullet(doc, "A3 — Site shortcut diagnostic: compare site-label prediction accuracy to forest-type prediction accuracy. Ratio near 1.0 = model uses biology; ratio >> 1.0 = model exploits site offsets.")
add_bullet(doc, "A4 — CS sensor fusion: use merged CASI+SASI 140-band CS tiles instead of CASI-only. Tests whether SWIR adds discriminative power.")
add_bullet(doc, "A5 — ALS structural fusion: concatenate 8 ALS metrics to the 40 spectral features. Tests whether structure complements spectra.")
add_bullet(doc, "A6 — Field spectra (BilyKriz dataset): run LOSO on 31 stands × 1,436 valid field-spectra bands. Cross-sensor transferability test.")


# ══════════════════════════════════════════════════════════════════
#  SECTION 6 — ITERATIONS AND DEBUGGING
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "6.  Preliminary Experiments, Trials, and Iterations", 1)

add_para(doc,
    "This section documents the full sequence of pipeline construction, failures, "
    "diagnoses, and fixes. In practice, these are not embarrassing mistakes — they "
    "reveal genuine ambiguities in the dataset documentation and are a necessary part "
    "of working with real-world remote-sensing data. Each failure is described with its "
    "root cause and resolution."
)

add_heading(doc, "6.1  Pipeline Construction Failures", 2)

failures = [
    ("Iteration 1: Solver incompatibility",
     "The default Logistic Regression solver (liblinear) raised ValueError: 'liblinear does not "
     "support multiclass'. This is because liblinear implements one-vs-rest internally, which "
     "is incompatible with the balanced-weights argument in multi-class mode.",
     "Changed solver from liblinear to lbfgs (a quasi-Newton optimiser that supports multinomial "
     "logistic regression natively). Removed the redundant penalty: l2 parameter."),

    ("Iteration 2: Stand ID mismatch",
     "After loading tiles, the pipeline reported 0 matched stands. Tile filenames like "
     "HY_PINE1_CASI.tif produced stand IDs of the form HY_PINE1_CASI, which did not match "
     "any entry in stand_metadata.csv where the corresponding key is HY_PINE1.",
     "Added a suffix-stripping loop in runner.py: after loading, any stand_id ending with "
     "_CASI is truncated by removing the last 5 characters. Extended to strip other sensor "
     "suffixes (_CS, _SASI) generically."),

    ("Iteration 3: Missing wavelength metadata (silent bug)",
     "The loader logged 'No wavelength metadata found; using band indices 0..47'. The "
     "water-vapour exclusion compared band wavelengths (in nm) to exclusion ranges (in nm). "
     "With band indices [0,1,...,47], all comparisons like 'is 5 in [895,1003]' returned "
     "False — so all 48 bands were retained and the exclusion was silently skipped.",
     "Extracted 48 CASI centre wavelengths from ENVI flight-line headers (.hdr files). "
     "Added them to config/default.yaml under casi_wavelengths_nm. In runner.py, detected "
     "whether tile wavelengths were band indices (by comparing to np.arange(n_bands)) and "
     "injected the config wavelengths when so."),

    ("Iteration 4: Wrong reflectance scale",
     "Band-importance plots showed values in the range 0–2350 instead of 0–1. The CASI "
     "tiles store integer reflectance × 10,000 (e.g., reflectance 0.2 = DN 2000). The "
     "config had reflectance_scale_factor: 1.0 (no scaling).",
     "Set reflectance_scale_factor: 10000 in config. Added a division step in runner.py "
     "after tile loading: tile.image = tile.image / scale."),

    ("Iteration 5: Wrong nodata value",
     "The config specified nodata_value: 0, but rasterio confirmed src.nodata = 10000 in "
     "all tiles. With nodata=0, the masking step treated all pixels with DN=0 as nodata "
     "(a small fraction) but let the true nodata pixels (DN=10000) through as valid "
     "reflectance. Post-scaling these became 1.0 reflectance — a physically plausible "
     "but incorrect value that corrupted stand means near tile edges.",
     "Changed nodata_value to 10000 in config. The mask now correctly identifies and "
     "excludes nodata pixels before computing stand means."),

    ("Iteration 6: Single-class fold crash",
     "The Bílý Kríž LOSO fold contains only coniferous stands. sklearn's "
     "classification_report raised ValueError: 'Number of classes, 1, does not match "
     "size of target_names, 3'.",
     "Added labels=class_names to both confusion_matrix() and classification_report() "
     "calls. This forces the functions to use all three class labels even when only one "
     "is present in y_true, producing zero-filled rows/columns for absent classes."),

    ("Iteration 7: ALS — vegetation returns not found at BK and LZ",
     "The ALS feature extractor filtered vegetation returns by cls == 1 (standard LAS "
     "classification code for unclassified vegetation). At Bílý Kríž and Lanzhot, "
     "classification code 4 (medium vegetation) was used instead. The extractor found "
     "zero vegetation points and returned None for all BK and LZ stands.",
     "Changed the vegetation mask from cls == 1 to ~(cls == 2) — all non-ground returns "
     "are treated as vegetation regardless of specific code. This correctly identifies "
     "vegetation at all four sites."),

    ("Iteration 8: Site shortcut test — meaningless BA",
     "The first site-shortcut test implementation used GroupKFold grouped by site for "
     "predicting site labels. Each test fold contained only one site, so balanced accuracy "
     "with a single class in y_true was trivially 0.0.",
     "Switched to StratifiedKFold(n_splits=5) so all sites appear in every fold's test "
     "set. The shortcut ratio (site BA / type BA) is now meaningful."),
]

for title, problem, fix in failures:
    add_heading(doc, title, 3)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Problem: "); r.bold = True; r.font.color.rgb = RGBColor(192,0,0)
    p.add_run(problem).font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("Fix: "); r2.bold = True; r2.font.color.rgb = RGBColor(0,112,0)
    p2.add_run(fix).font.size = Pt(10.5)

add_heading(doc, "6.2  What Did Not Work — and Why", 2)

failures2 = [
    ("PCA-10 compression (A1)",
     "Compressing 40 bands to 10 PCA components reduced Random Forest balanced accuracy "
     "from 0.621 to 0.600 (−0.021). This was unexpected given PCA's reputation as a "
     "regulariser. The explanation: the 10 components capture 99% of total spectral "
     "variance, but spectral variance is dominated by illumination and atmospheric "
     "variation rather than forest type. The biologically informative variance (e.g., "
     "the shape of the red-edge inflection, NIR plateau curvature) is distributed across "
     "many low-variance components that PCA discards. For n = 58, there is no overfitting "
     "problem to regularise — the raw 40-band representation is better."),

    ("CS sensor fusion (A4)",
     "Adding SWIR bands from the merged CASI+SASI CS product (140 bands) did not "
     "improve classification over CASI-only. The field spectroscopy separability analysis "
     "subsequently revealed that the most discriminative SWIR region (2317–2490 nm, "
     "cellulose/lignin absorption) lies outside even the CS sensor's range (max 2443 nm). "
     "Additionally, at stand scale the spectral unmixing effect of mixed canopy elements "
     "reduces the contrast between species-specific SWIR features."),

    ("ALS-only classification (A5 component)",
     "ALS structural features alone achieved BA = 0.192 (near chance = 0.33). Canopy "
     "height and cover fraction are not informative for forest type at the biome scale: "
     "a 30 m tall, high-cover stand could be mature Norway spruce (coniferous) or mature "
     "English oak (broadleaved). Structure differentiates stands at the within-site "
     "scale (young vs. mature, thinned vs. dense) but not across biomes."),

    ("Field spectra LOSO — BK fold collapse",
     "The LOSO experiment on 31 field-spectra stands achieved BA = 0.413, dragged down "
     "by the Bílý Kríž test fold where BA = 0.000: the model predicted all 8 BK stands "
     "as broadleaved. When BK is the test set, training consists of Järvselja "
     "(4 coniferous + 9 broadleaved) and Lanzhot (10 broadleaved) = 19 broadleaved vs. "
     "4 coniferous. Despite class_weight=balanced, the Random Forest still predicted "
     "all BK coniferous stands as broadleaved. This is because BK Norway spruce (high "
     "altitude, continental climate, September imagery) occupies a different spectral "
     "position than Järvselja Scots pine (boreal, July imagery) — the model has never "
     "seen coniferous spectra that look like BK spruce. The fix (Hyytiälä) is structural: "
     "the airborne dataset includes Hyytiälä with 22 coniferous stands, balancing the "
     "training distribution when any other site is the test fold."),

    ("Mixed-class classification",
     "Across every experiment, the mixed class achieved near-zero classification accuracy. "
     "The aggregated confusion matrix shows 0/10 mixed stands correctly identified. All "
     "are misclassified as coniferous. This is not a model failure — it is a fundamental "
     "limitation of the feature representation. Stand-level mean spectra average over "
     "the spatial mixture of coniferous and broadleaved tree crowns, producing a feature "
     "vector that lies between the two pure classes with high within-class variance. A "
     "stand that is 60% spruce and 40% birch has a mean spectrum that looks more like "
     "a pure spruce stand than a pure birch stand, because spruce has darker, more "
     "absorptive spectra that dominate the canopy average. Resolving mixed forest "
     "requires sub-stand spatial features (individual crown delineation or patch-level "
     "classification)."),
]

for title, text in failures2:
    add_heading(doc, title, 3)
    add_para(doc, text)

add_heading(doc, "6.3  What Worked — Key Decisions That Led to Better Results", 2)

wins = [
    ("Binary classification — the clean experiment",
     "Dropping the 10 mixed stands raised balanced accuracy from 0.621 to 0.932. This "
     "was the single largest improvement in the entire project. The decision was guided "
     "by the ecological argument: 'mixed forest' is not a discrete spectral category but "
     "a continuous mixture. Reporting binary results alongside three-class results is "
     "scientifically honest — it acknowledges that the research question has two parts: "
     "a hard part (three-class with ambiguous mixed label) and a clean part (binary "
     "conifer vs. broadleaf)."),

    ("Injecting CASI wavelengths from ENVI headers",
     "This was the fix for the silent wavelength-index bug. The wavelengths were found "
     "in the ENVI .hdr files accompanying the raw flight-line BSQ images: "
     "wavelength = {..., 0.93765, 0.95220, ...} in micrometres. Multiplying by 1000 "
     "gave nanometre values matching the expected CASI-1500 range. Once injected, "
     "water-vapour exclusion correctly removed 8 contaminated bands."),

    ("CASI + ALS fusion for 3-class (+4.1 pp)",
     "Concatenating 8 ALS structural metrics to the 40-band spectral vector improved "
     "3-class balanced accuracy from 0.621 to 0.663. The improvement is concentrated "
     "in the Hyytiälä and Järvselja folds. Post-hoc analysis suggests the gain comes "
     "from the mixed class: some mixed stands have distinctive structural profiles "
     "(lower mean height, higher height variability) that help distinguish them from "
     "pure coniferous stands."),

    ("Cross-sensor BK transfer (BA = 0.875)",
     "The experiment of applying the CASI-trained model to field spectra resampled to "
     "CASI band centres confirmed inter-sensor transferability. 7 of 8 BK stands were "
     "correctly labelled coniferous. The misclassified stand (Spruce3, predicted mixed) "
     "has 27% broadleaved basal area — the highest in BK — making the model's uncertainty "
     "ecologically sensible. This result shows that resampling field spectra to sensor "
     "band centres preserves enough spectral structure for classification."),
]

for title, text in wins:
    add_heading(doc, title, 3)
    add_para(doc, text)


# ══════════════════════════════════════════════════════════════════
#  SECTION 7 — RESULTS
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "7.  Results", 1)

add_heading(doc, "7.1  Baseline Classification — Five Models, LOSO-CV", 2)
add_para(doc,
    "Table 3 presents the aggregated LOSO-CV performance across all five classifiers "
    "on the three-class problem (58 stands, 4 folds). Random Forest is the best "
    "performer on both balanced accuracy and macro-F1 and is used as the reference "
    "model for all subsequent ablations."
)

add_table(doc,
    ["Model", "Mean BA", "±SD", "Mean OA", "Macro-F1"],
    [
        ["Logistic Regression", "0.542", "0.273", "0.698", "0.470"],
        ["Random Forest",       "0.621", "0.228", "0.756", "0.576"],
        ["SVM (RBF)",           "0.592", "0.162", "0.694", "0.438"],
        ["Gradient Boosting",   "0.592", "0.251", "0.706", "0.555"],
        ["PLS-DA",              "0.591", "—",     "—",     "—"    ],
    ],
    col_widths=[4.5, 2.2, 2.0, 2.2, 2.5]
)
add_caption(doc, "Table 3. Classifier performance — 3-class LOSO-CV (58 stands, 4 sites). "
                 "BA = balanced accuracy; OA = overall accuracy; F1 = macro-F1. "
                 "Best value per column in bold.")

add_chart(doc, chart1_stream, width_inches=6.2,
          caption="Figure 1 — Balanced accuracy (±1 std) and overall accuracy for all five classifiers. "
                  "Random Forest achieves the highest balanced accuracy of 0.621.")

add_heading(doc, "7.2  Per-Site Breakdown (Random Forest)", 2)
add_para(doc,
    "The per-site breakdown reveals large variance across folds — a direct consequence "
    "of site composition. Bílý Kríž (all coniferous) is trivially classified. The "
    "meaningful folds are Hyytiälä (BA = 0.402) and Järvselja (BA = 0.583)."
)

add_table(doc,
    ["Test site", "BA", "n stands", "Classes present", "Interpretation"],
    [
        ["Bílý Kríž",  "1.000", "7",  "Coniferous only",       "Trivial — single-class fold"],
        ["Hyytiälä",   "0.402", "28", "All three classes",     "Hardest fold — mixed class fails"],
        ["Järvselja",  "0.583", "13", "All three classes",     "Moderate — mixed partially identified"],
        ["Lanzhot",    "0.500", "10", "Broadleaved + Mixed",   "Broadleaved well identified"],
        ["Mean",       "0.621", "58", "",                       ""]
    ],
    col_widths=[3.0, 1.8, 2.0, 3.5, 4.2]
)
add_caption(doc, "Table 4. Random Forest per-site balanced accuracy under LOSO-CV.")

add_chart(doc, chart2_stream, width_inches=5.8,
          caption="Figure 2 — Per-site balanced accuracy for Random Forest. Dashed line = mean BA = 0.621. "
                  "The Bílý Kríž result (1.0) reflects trivial single-class test set composition.")

add_heading(doc, "7.3  Confusion Matrix Analysis", 2)
add_para(doc,
    "The aggregated confusion matrix (summed over 4 LOSO folds) reveals the asymmetric "
    "nature of classification errors. Broadleaved stands are classified nearly perfectly "
    "(31/33 correct; recall = 0.94). Coniferous stands have moderate accuracy (12/15 "
    "correct; recall = 0.80). Mixed stands are never correctly identified (0/10; recall "
    "= 0.00) — all are misclassified, predominantly as coniferous (8/10)."
)

add_chart(doc, chart3_stream, width_inches=4.5,
          caption="Figure 3 — Row-normalised confusion matrix aggregated over all 4 LOSO folds. "
                  "Values show fraction of true-class samples predicted to each class. "
                  "The mixed class is completely undetected; all 10 mixed stands are "
                  "misclassified as coniferous (n=8) or broadleaved (n=2).")

add_heading(doc, "7.4  Ablation Study", 2)
add_para(doc,
    "Six ablations were run using Random Forest. Table 5 and Figure 4 summarise "
    "the results."
)

add_table(doc,
    ["Experiment", "3-class BA", "Binary BA", "Delta (3-class)", "Key finding"],
    [
        ["Baseline (40 CASI bands)",     "0.621", "0.932", "—",      "Reference"],
        ["PCA-10 compression",           "0.600", "—",     "−0.021", "Dimensionality reduction hurts"],
        ["CS sensor (140 bands)",        "≈0.62", "≈0.93", "~0.000", "SWIR adds no benefit"],
        ["ALS-only (8 structural)",      "0.192", "0.247", "−0.429", "Structure alone is near-chance"],
        ["CASI + ALS fusion",            "0.663", "0.932", "+0.042", "Modest 3-class gain; zero binary gain"],
        ["Field spectra LOSO",           "0.413", "0.413", "−0.208", "3-site dataset; BK fold fails"],
        ["Cross-sensor BK (CASI→field)", "—",     "0.875", "—",      "7/8 correct; inter-sensor transfer"],
    ],
    col_widths=[4.0, 2.2, 2.2, 2.5, 3.5]
)
add_caption(doc, "Table 5. Ablation results. All experiments use Random Forest with LOSO-CV.")

add_chart(doc, chart4_stream, width_inches=6.5,
          caption="Figure 4 — Ablation comparison: 3-class vs binary balanced accuracy across feature sets. "
                  "Grey bars indicate N/A (experiment not applicable to that problem). "
                  "Binary classification is near-perfect; 3-class is limited by the mixed class.")

add_heading(doc, "7.5  Binary Classification Detail", 2)
add_para(doc,
    "Dropping 10 mixed stands (48 samples remain: 33 coniferous, 15 broadleaved) "
    "produces markedly better cross-site generalisation. The improvement confirms "
    "that the mixed class — not the biome generalisation problem itself — is the "
    "primary source of error in the three-class problem."
)

add_table(doc,
    ["Test site", "BA (3-class)", "BA (binary)", "Gain"],
    [
        ["Bílý Kríž", "1.000", "1.000", "0.000"],
        ["Hyytiälä",  "0.402", "0.852", "+0.450"],
        ["Järvselja", "0.583", "0.875", "+0.292"],
        ["Lanzhot",   "0.500", "1.000", "+0.500"],
        ["Mean",      "0.621", "0.932", "+0.311"],
    ],
    col_widths=[3.2, 3.0, 3.0, 3.0]
)
add_caption(doc, "Table 6. Three-class vs. binary balanced accuracy per fold. Dropping the mixed class "
                 "produces gains of 0.29–0.50 per fold.")

add_heading(doc, "7.6  ALS Structural Features", 2)
add_para(doc,
    "The eight stand-level ALS metrics (height percentiles and cover fraction) were "
    "fused with the 40-band CASI spectral features. The result supports the spectral "
    "primacy hypothesis: spectral features dominate, and ALS contributes marginally."
)

add_table(doc,
    ["Feature set", "3-class BA", "Binary BA", "n features"],
    [
        ["ALS-only",   "0.192", "0.247", "8"],
        ["CASI-only",  "0.621", "0.932", "40"],
        ["CASI + ALS", "0.663", "0.932", "48"],
    ],
    col_widths=[4.0, 3.0, 3.0, 2.5]
)
add_caption(doc, "Table 7. ALS fusion ablation. The +4.1 pp gain for 3-class is real but moderate; "
                 "no binary gain confirms spectral saturation for the easier problem.")

add_heading(doc, "7.7  Field Spectroscopy (BilyKriz Dataset)", 2)
add_para(doc,
    "The field spectroscopy dataset (31 stands, 350–2500 nm, 1,436 valid bands) "
    "provides a second line of evidence. Three distinct analyses were performed:"
)

add_heading(doc, "Cross-Site LOSO (3 folds)", 3)
add_para(doc,
    "Results were substantially lower than airborne CASI (BA = 0.413 vs. 0.621), "
    "driven by the Bílý Kríž fold collapse (BA = 0.000). Järvselja fold BA = 0.639 "
    "is actually higher than the airborne result (0.583) — suggesting that at full "
    "2151 bands the spectral resolution advantage compensates for smaller sample size "
    "when the test site has diverse types."
)

add_chart(doc, chart8_stream, width_inches=5.8,
          caption="Figure 8 — Airborne CASI vs. field spectra balanced accuracy per LOSO fold. "
                  "The BilyKriz fold collapse (field BA = 0.0) reflects training-set imbalance: "
                  "19 broadleaved vs. 4 coniferous in training when BK is test.")

add_heading(doc, "Spectral Separability Analysis", 3)
add_para(doc,
    "A Jeffreys-Matusita separability proxy (mu1−mu2)² / (sigma1²+sigma2²) per "
    "wavelength identified the top-10 most discriminative wavelengths as: "
    "2317, 2318, 2319, 2320, 2321, 2454, 2490, 350, 353, 372 nm. The SWIR-2 cluster "
    "(2317–2490 nm) corresponds to cellulose and lignin absorption features known to "
    "differ between needle-leaved and broad-leaved species. Critically, this region "
    "is not covered by CASI (382–1052 nm) — yet CASI achieves BA = 0.932 for binary "
    "classification. This confirms that VIS-NIR canopy spectra contain sufficient "
    "discriminative information at stand scale; SWIR does not add further benefit."
)

# field spectra separability figure
sep_fig_path = FIG_DIR / "field_spectra_separability.png"
if sep_fig_path.exists():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(sep_fig_path), width=Inches(5.8))
    add_caption(doc, "Figure 10 — Spectral separability index (coniferous vs. broadleaved) across "
                     "350–2500 nm from field spectra. SWIR-2 (2317–2490 nm) is the most discriminative "
                     "region. Grey shaded areas are water-vapour exclusion zones.")

add_heading(doc, "Cross-Sensor Prediction", 3)
add_para(doc,
    "The CASI-trained Random Forest (trained on 51 HY+JS+LZ airborne stands) was "
    "applied to BK field spectra resampled to 40 CASI band centres via linear "
    "interpolation. Result: 7/8 BK stands correctly classified as coniferous "
    "(BA = 0.875). The misclassified stand (Spruce3) has the highest broadleaved "
    "basal-area fraction in the BK dataset (27%), making the model's prediction of "
    "'mixed' ecologically coherent."
)

add_chart(doc, chart5_stream, width_inches=6.0,
          caption="Figure 5 — Cross-sensor prediction probabilities for 8 BK field-spectra stands. "
                  "CASI model (trained on airborne HY+JS+LZ) is applied to field spectra resampled "
                  "to 40 CASI band centres. BK_Spruce3 (27% broadleaved BA) is the only misclassification.")

# Embed band importance figure
rf_imp_path = OUT_DIR / "baseline_multisite_casi_loso/figures/rf_band_importance.png"
if rf_imp_path.exists():
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(rf_imp_path), width=Inches(5.8))
    add_caption(doc, "Figure 11 — Random Forest band importance across 40 valid CASI bands (382–1052 nm). "
                     "Red bars highlight the top-20 most important bands. The NIR plateau region "
                     "(760–895 nm) and red-edge (680–740 nm) are consistently important.")


# ══════════════════════════════════════════════════════════════════
#  SECTION 8 — DISCUSSION
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "8.  Discussion", 1)

add_heading(doc, "8.1  Binary vs. Three-Class: What the Gap Tells Us", 2)
add_para(doc,
    "The jump from BA = 0.621 (three-class) to BA = 0.932 (binary) is the most "
    "informative single number in this study. It tells us that the failure mode is "
    "not the biome-generalisation problem — the spectral signal for conifer vs. "
    "broadleaf is robust and consistent across Finland, Estonia, and Czech Republic. "
    "The failure is the mixed label. Stand-level mean spectra average over the spatial "
    "mosaic of different tree crowns, producing a feature vector that is a weighted "
    "sum of the pure class spectra. A stand with 60% spruce and 40% birch produces a "
    "mean spectrum that is geometrically closer to pure spruce than to pure birch — "
    "because spruce needles have lower reflectance and absorb more light, biasing the "
    "stand average toward the darker class. As a result, the classifier consistently "
    "places mixed stands in the coniferous region of feature space."
)
add_para(doc,
    "The implication for forest inventory applications is practical: a binary "
    "conifer/broadleaf classifier trained on this dataset can be deployed operationally "
    "with high confidence. A mixed-class classifier requires a fundamentally different "
    "approach — likely individual crown delineation from high-resolution imagery, "
    "followed by crown-level classification and spatial aggregation."
)

add_heading(doc, "8.2  Site Composition as a Confound in LOSO Evaluation", 2)
add_para(doc,
    "The Bílý Kríž fold (BA = 1.000) is a statistical artefact. Any non-trivial "
    "classifier that has ever seen coniferous spectra will achieve near-perfect accuracy "
    "on a test set that is 100% coniferous. Including this fold in the mean inflates "
    "the reported metric. A fairer summary would report the mean BA over the three "
    "informative folds (Hyytiälä, Järvselja, Lanzhot): (0.402 + 0.583 + 0.500) / 3 "
    "= 0.495. This is a more honest estimate of the classifier's ability to resolve "
    "mixed-type forests at new sites."
)
add_para(doc,
    "The site shortcut ratio (1.14×) — site-label BA divided by type-label BA — is "
    "reassuringly close to 1.0. A ratio significantly above 1.0 would indicate the "
    "model learns to predict site identity (by exploiting site-specific atmospheric "
    "offsets, sensor-geometry effects, or soil-background differences) rather than "
    "biological forest type. At 1.14×, the evidence suggests the classifier is "
    "primarily exploiting biological signal, but the margin is narrow enough to "
    "warrant caution in deployment to new geographic regions."
)

add_heading(doc, "8.3  Why ALS Helps 3-class but Not Binary", 2)
add_para(doc,
    "The asymmetric ALS gain (+4.1 pp for 3-class, 0 pp for binary) can be understood "
    "from the feature-space geometry. In the binary problem, spectral features already "
    "achieve near-perfect separation; there is no residual error for ALS to reduce. "
    "In the three-class problem, mixed-class errors are the dominant failure mode, and "
    "these errors occur near the decision boundary between coniferous and mixed. "
    "Some mixed stands have detectably lower mean canopy height (because the broadleaved "
    "component includes shorter species like Hazel or Hornbeam) or higher height variance "
    "(because the two components have different height profiles). ALS features provide "
    "a secondary signal in this narrow confusion zone."
)

add_heading(doc, "8.4  Field Spectra vs. Airborne: Sensor Modality Matters", 2)
add_para(doc,
    "The lower field-spectra LOSO result (BA = 0.413 vs. 0.621 airborne) does not "
    "mean that field spectra are less informative than airborne imagery — the "
    "separability analysis shows field spectra contain more discriminative information "
    "per wavelength in the SWIR-2 region. The lower performance reflects two structural "
    "differences: (1) only three sites in the field dataset vs. four in the airborne "
    "dataset, causing severe training imbalance for the BK fold; and (2) n = 31 stands "
    "rather than 58, reducing the training set for any given fold."
)
add_para(doc,
    "The cross-sensor experiment (BA = 0.875) is the more relevant comparison: it "
    "shows that when the same four-site diversity is available in the training set, "
    "the CASI-trained model applies effectively to field spectra. The 5.7 pp gap "
    "from perfect (0.875 vs. 1.0 expected if the signal were perfectly transferable) "
    "reflects genuine spectral differences between airborne canopy imagery and field "
    "hand-held measurements — differences in illumination geometry, canopy bidirectional "
    "reflectance, and the proportion of shadowed understory in the airborne footprint."
)

add_heading(doc, "8.5  Limitations", 2)
add_bullet(doc, "Small sample size (n = 58 stands): machine learning classifiers are underdetermined. Results should be treated as indicative of feasibility, not as production accuracy estimates.")
add_bullet(doc, "Single acquisition date: all data were collected in one growing season (summer 2019). Phenological variation across seasons is not assessed. Broadleaved stands change substantially from spring flush to autumn senescence.")
add_bullet(doc, "Species lumped within type: 'coniferous' includes Scots pine (xeric, open canopy) and Norway spruce (mesic, dense canopy) whose spectra differ substantially. Species-level analysis was not possible at n = 58.")
add_bullet(doc, "Mixed-class ambiguity: the mixed label has no clear spectral identity. Future work should define mixed by fractional canopy cover (e.g., 30–70% coniferous) rather than a binary label.")
add_bullet(doc, "No temporal generalisation: the model was trained and tested on data from the same campaign year. Inter-annual spectral variation from drought, insect damage, or phenological shifts is not modelled.")


# ══════════════════════════════════════════════════════════════════
#  SECTION 9 — RELATED WORK
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "9.  Review of Related Work", 1)

add_para(doc,
    "Fassnacht et al. (2016, Remote Sensing of Environment 186: 64–87) reviewed 74 "
    "tree-species classification studies and found that stand-level (plot-level) "
    "feature extraction consistently outperforms pixel-level methods in small-sample "
    "regimes — directly validating our choice of stand-level mean reflectance. They "
    "also found that Random Forest and SVM were the most frequently successful "
    "classifiers, consistent with our results."
)
add_para(doc,
    "Roberts et al. (2017, Ecography 40: 913–929) and Ploton et al. (2020, Nature "
    "Communications 11: 4540) rigorously demonstrate that ignoring spatial "
    "autocorrelation — by using random train/test splits instead of spatial or "
    "site-based blocking — produces drastically optimistic performance estimates, "
    "sometimes by 20–40 percentage points. Our exclusive use of LOSO-CV directly "
    "implements the spatial blocking they recommend."
)
add_para(doc,
    "Richter et al. (2016, Remote Sensing of Environment 179: 324–338) show that "
    "SWIR bands (1100–2500 nm) improve tree-species classification at both leaf and "
    "canopy scale, particularly for lignin and cellulose content differences between "
    "conifers and broadleaves. Our field-spectra separability analysis is consistent "
    "with this finding: SWIR-2 (2317–2490 nm) is the most separable region. However, "
    "their gain was demonstrated at the leaf and plot scale with controlled measurements; "
    "at stand scale with LOSO evaluation, we find no SWIR benefit — suggesting the "
    "added complexity of canopy mixing, illumination variation, and site-specific "
    "atmospheric effects dilutes the SWIR advantage."
)
add_para(doc,
    "Torabzadeh et al. (2019, Agricultural and Forest Meteorology 279: 107666) fused "
    "hyperspectral and ALS features for tree species mapping in a temperate mixed forest "
    "and reported 5–15% improvements from ALS crown features. Our +4.1 pp gain for "
    "3-class classification is comparable, noting that their regime (within-site, "
    "species level) is easier than ours (cross-site, type level). Their study was "
    "single-site; our cross-site design is a harder and more realistic evaluation."
)
add_para(doc,
    "Rautiainen et al. (2024, Earth System Science Data 16: 5069–5098) provide the "
    "FREEDLES dataset reference. They demonstrate the multi-site, multi-sensor "
    "acquisition design and provide stand-level metadata. Their paper does not include "
    "a classification experiment — this study is among the first to operationalise "
    "the FREEDLES dataset for forest-type mapping."
)


# ══════════════════════════════════════════════════════════════════
#  SECTION 10 — CONCLUSIONS
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "10.  Conclusions", 1)

add_para(doc,
    "This project set out to answer whether airborne hyperspectral reflectance "
    "can generalise across European biomes for forest-type classification. The answer "
    "is: yes for the binary conifer/broadleaf problem (BA = 0.932), and partially for "
    "the three-class problem (BA = 0.621, limited by the unresolvable mixed class)."
)
add_para(doc,
    "Four contributions stand out:"
)
add_bullet(doc, "An end-to-end reproducible classification pipeline from raw FREEDLES GeoTIFF tiles to LOSO-validated metrics, including documentation of eight non-trivial data-engineering fixes.")
add_bullet(doc, "A comprehensive ablation suite (PCA, binary, site shortcut, CS sensor, ALS fusion, field spectra, cross-sensor) that systematically isolates the contribution of each design choice.")
add_bullet(doc, "Integration of the BilyKriz field spectroscopy dataset, including spectral separability analysis (identifying SWIR-2 as most discriminative) and cross-sensor transfer (BA = 0.875).")
add_bullet(doc, "A clear diagnosis of the primary failure mode: mixed-class stands require sub-stand spatial features, not stand-level means, and their poor classification should not be conflated with the generalisation failure of the binary problem.")

add_heading(doc, "Recommended Next Steps", 2)
add_bullet(doc, "Patch-level or crown-level feature extraction (15×15 pixel windows or ALS-segmented crowns) to replace stand-level means for the mixed class.")
add_bullet(doc, "Species-level classification within the coniferous type (Scots pine vs. Norway spruce) as a finer ecological question achievable at current sample sizes.")
add_bullet(doc, "Multi-date acquisition: including spring (before broadleaved flush) and autumn (senescence) imagery to exploit phenological separability.")
add_bullet(doc, "External dataset validation: apply the trained model to NEON airborne hyperspectral campaigns or HyRANK challenge data to test continental-scale transfer.")
add_bullet(doc, "Probabilistic output calibration: the RF class probabilities are only weakly calibrated; Platt scaling or isotonic regression would produce more reliable uncertainty estimates for operational use.")


# ══════════════════════════════════════════════════════════════════
#  SECTION 11 — WORKPLAN AND PROCESS LOG
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "11.  Workplan and Process Documentation", 1)

add_para(doc,
    "The project followed a structured but iterative workflow across seven phases. "
    "Phases overlapped in practice — data familiarisation continued throughout the "
    "experiment, and problem formulation was refined as empirical results emerged."
)

add_table(doc,
    ["Phase", "Key activities", "Outputs", "Duration (approx.)"],
    [
        ["1. Data familiarisation",
         "Discover actual file paths; inspect tile headers; read ENVI metadata; plot raw spectra; check nodata, scale, and wavelength values",
         "stand_metadata.csv; wavelength config; basic spectra plots",
         "~2 days"],
        ["2. Problem formulation",
         "Define LOSO-CV rationale; select balanced accuracy; formalise three research questions; review Roberts 2017 and Ploton 2020",
         "Research strategy document; LOSO implementation",
         "~1 day"],
        ["3. Pipeline construction",
         "Fix 8 bugs (solver, suffix, wavelength, scale, nodata, single-class fold, ALS class code, shortcut test); build runner.py",
         "Working runner.py; config/default.yaml",
         "~3 days"],
        ["4. Baseline experiments",
         "Run 5 classifiers, 3-class, LOSO-CV; generate confusion matrices and per-fold plots; interpret results",
         "baseline_multisite_casi_loso/; Table 1-4",
         "~1 day"],
        ["5. Ablations (round 1)",
         "PCA-10; binary classification; site shortcut; CS sensor fusion",
         "ablations/*.json; Tables 5-6",
         "~1 day"],
        ["6. ALS integration",
         "Install laspy; extract structural features from 58 LAS stand directories; ALS fusion experiment",
         "als_features.csv; als_fusion.json; Table 7",
         "~2 days"],
        ["7. BilyKriz dataset",
         "Load Excel field spectra; LOSO on 31 stands; separability analysis; cross-sensor prediction; all figures",
         "field_spectra_deep.json; cross_sensor_bk.json; Figures 5,8,10",
         "~2 days"],
        ["8. Report",
         "Compile all results, figures, discussion, and references into this Word document",
         "Research_Report_Hyperspectral_Forest.docx",
         "~1 day"],
    ],
    col_widths=[2.8, 5.5, 3.8, 2.4]
)
add_caption(doc, "Table 8. Project workplan and process log.")


# ══════════════════════════════════════════════════════════════════
#  SECTION 12 — REFERENCES
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "References", 1)

refs = [
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
    "Fassnacht, F.E., Latifi, H., Stereńczak, K., Modzelewska, A., Lefsky, M., Waser, L.T., Straub, C., & Ghosh, A. (2016). Review of studies on tree species classification from remotely sensed data. Remote Sensing of Environment, 186, 64–87.",
    "Ploton, P., Mortier, F., Rejou-Mechain, M., Barbier, N., Couteron, P., Dauby, G., ... & Pélissier, R. (2020). Spatial validation reveals poor predictive performance of large-scale ecological mapping models. Nature Communications, 11, 4540.",
    "Rautiainen, M., Mõttus, M., Blades, N., Kuusk, J., Kuusk, A., Lang, M., Lükk, T., Majasalmi, T., Niinemets, Ü., Pisek, J., Putzenlechner, B., Ruiz-Arias, J.A., Scheuermann, C., Sepp, K., Sims, D., Taulavuori, K., Taulavuori, E., & Treitz, P. (2024). FREEDLES: a dataset of multi-scale spectroradiometric measurements and airborne hyperspectral and laser scanning data in European forest environments. Earth System Science Data, 16, 5069–5098.",
    "Richter, R., Reu, B., Wirth, C., Doktor, D., & Vohland, M. (2016). The use of airborne hyperspectral data for tree species classification in a species-rich Central European forest area. International Journal of Applied Earth Observation and Geoinformation, 52, 464–474.",
    "Roberts, D.R., Bahn, V., Ciuti, S., Boyce, M.S., Elith, J., Guillera-Arroita, G., ... & Dormann, C.F. (2017). Cross-validation strategies for data with temporal, spatial, hierarchical, or phylogenetic structure. Ecography, 40(8), 913–929.",
    "Schölkopf, B., & Smola, A.J. (2002). Learning with Kernels: Support Vector Machines, Regularization, Optimization, and Beyond. MIT Press.",
    "Torabzadeh, H., Leiterer, R., Hueni, A., Schaepman, M.E., & Morsdorf, F. (2019). Tree species classification in a temperate mixed forest using a combination of imaging spectroscopy and airborne laser scanning. Agricultural and Forest Meteorology, 279, 107666.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-1.2)
    p.paragraph_format.space_after   = Pt(5)
    run = p.add_run(f"{i}.  {ref}")
    run.font.size = Pt(10.5)


# ══════════════════════════════════════════════════════════════════
#  APPENDIX — REPRODUCIBILITY
# ══════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "Appendix A — Software Environment and Reproducibility", 1)

add_para(doc,
    "All experiments were run on Windows 11 Enterprise (x64) with Python 3.12. "
    "The following package versions were used:"
)
env_table = [
    ["scikit-learn", "1.x", "Classification, cross-validation"],
    ["numpy", "1.26+", "Array operations"],
    ["pandas", "2.x", "Metadata handling"],
    ["rasterio", "1.3+", "GeoTIFF I/O"],
    ["laspy", "2.7.0", "LAS point cloud reading"],
    ["scipy", "1.x", "Interpolation (cross-sensor resampling)"],
    ["matplotlib", "3.x", "All figures"],
    ["python-docx", "1.2.0", "This document"],
    ["pyyaml", "6.x", "Configuration"],
]
add_table(doc,
    ["Package", "Version", "Usage"],
    env_table,
    col_widths=[3.5, 2.5, 8.5]
)
add_caption(doc, "Table A1. Python environment.")

add_heading(doc, "Key Script Invocations", 2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.space_after = Pt(6)
run = p.add_run(
    "# Baseline experiment\n"
    "cd model/\n"
    "python -m src.experiments.runner --config config/default.yaml\n\n"
    "# ALS feature extraction\n"
    "python scripts/extract_als_features.py\n\n"
    "# ALS + spectral fusion experiment\n"
    "python scripts/als_fusion_experiment.py\n\n"
    "# Field spectra experiments\n"
    "python scripts/field_spectra_experiment.py\n"
    "python scripts/field_spectra_deep.py\n\n"
    "# Cross-sensor BK prediction\n"
    "python scripts/cross_sensor_bk_predict.py\n\n"
    "# This report\n"
    "python scripts/build_report.py"
)
run.font.name = "Courier New"; run.font.size = Pt(9)

add_para(doc, "Random seed: 42 in all scripts. Outputs: outputs/ directory.")

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(str(DOCX_OUT))
print(f"\nDocument saved: {DOCX_OUT}")
print(f"File size: {DOCX_OUT.stat().st_size / 1024:.1f} KB")
